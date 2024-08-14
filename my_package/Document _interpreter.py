import os
import docx
import PyPDF2
import csv
from openpyxl import load_workbook
from openpyxl import Workbook
from collections import Counter
import re


# 定义文件类型常量
FILE_TYPE_TEXT = 'text'
FILE_TYPE_PDF = 'pdf'
FILE_TYPE_DOCX = 'docx'
FILE_TYPE_CSV = 'csv'
FILE_TYPE_XLSX = 'xlsx'

def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except FileNotFoundError:
        print(f"文件 '{file_path}' 未找到.")
    except PermissionError:
        print(f"没有权限访问文件 '{file_path}'.")            
    except Exception as e:
        print(f"读取文件 '{file_path}' 时出现错误: {e}")
    return None

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            num_pages = reader.numPages
            text = ''
            for page_num in range(num_pages):
                page = reader.getPage(page_num)
                text += page.extract_text()
            return text
    except FileNotFoundError:
        print(f"文件 '{file_path}' 未找到.")
    except Exception as e:
        print(f"从文件 '{file_path}' 提取文本时出现错误: {e}")
    return None

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except FileNotFoundError:
        print(f"文件 '{file_path}' 未找到.")
    except Exception as e:
        print(f"读取文件 '{file_path}' 时出现错误: {e}")
    return None

def read_csv(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            lines = [','.join(row) for row in reader]
            return '\n'.join(lines)
    except FileNotFoundError:
        print(f"文件 '{file_path}' 未找到.")
    except Exception as e:
        print(f"读取文件 '{file_path}' 时出现错误: {e}")
    return None


def read_xlsx(file_path):
    try:
        wb = load_workbook(file_path)
        ws = wb.active
        data = []
        for row in ws.iter_rows():
            row_data = []
            for cell in row:
                row_data.append(str(cell.value))
            data.append(','.join(row_data))
        return '\n'.join(data)
    except FileNotFoundError:
        print(f"文件 '{file_path}' 未找到.")
    except Exception as e:
        print(f"读取文件 '{file_path}' 时出现错误: {e}")
    return None

def count_word_frequency(text):
    words = re.findall(r'\b\w+\b', text.lower())
    return Counter(words)

def extract_keywords(text, num_keywords=5):
    words = re.findall(r'\b\w+\b', text.lower())
    word_counts = Counter(words)
    return word_counts.most_common(num_keywords)

def search_and_replace(text, search_term, replace_term):
    return text.replace(search_term, replace_term)

def extract_table_data_from_pdf(file_path):
    try:
        table_data = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            num_pages = reader.numPages
            for page_num in range(num_pages):
                page = reader.getPage(page_num)
                text = page.extract_text()
                # 使用正则表达式来匹配表格数据的简单示例
                tables = re.findall(r'((?:[\w\s,]+;?)+)', text)
                for table in tables:
                    rows = table.split(';')
                    table_data.append(rows)
        return table_data
    except Exception as e:
        print(f"从文件 '{file_path}' 提取表格数据时出现错误: {e}")
    return None
    
def extract_specific_pages(input_file, output_file, page_numbers):
    try:
        with open(input_file, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            writer = PyPDF2.PdfFileWriter()
            for page_num in page_numbers:
                writer.addPage(reader.getPage(page_num - 1))  # Page numbers are 1-based
            with open(output_file, 'wb') as out_f:
                writer.write(out_f)
        print(f"成功从 '{input_file}' 提取指定页面到 '{output_file}'.")
    except Exception as e:
        print(f"提取PDF页面时出现错误: {e}")
        
def update_docx(file_path, new_content):
    try:
        # 为了安全，先备份文件
        backup_file = file_path + '.bak'
        shutil.copyfile(file_path, backup_file)        
        doc = docx.Document(file_path)
        # Example: Update the first paragraph with new content
        if doc.paragraphs:
            doc.paragraphs[0].text = new_content
            doc.save(file_path)
            print(f"成功更新 '{file_path}'.")
    except Exception as e:
        print(f"更新 '{file_path}' 时出现错误: {e}")

def clean_csv(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        with open(file_path, 'w', encoding='utf-8') as f:
            f.writelines(cleaned_lines)
        print(f"成功清洗 '{file_path}'.")
    except Exception as e:
        print(f"清洗 '{file_path}' 时出现错误: {e}")
        
def export_to_excel(data, output_file):
    try:
        wb = load_workbook()
        ws = wb.active
        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, value in enumerate(row_data, start=1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        wb.save(output_file)
        print(f"成功导出数据到 '{output_file}'.")
    except Exception as e:
        print(f"导出数据到 '{output_file}' 时出现错误: {e}")

def interpret_file(file_path):
    file_type = get_file_type(file_path)
    if file_type == FILE_TYPE_TEXT:
        return read_text_file(file_path)
    elif file_type == FILE_TYPE_PDF:
        return extract_text_from_pdf(file_path)
    elif file_type == FILE_TYPE_DOCX:
        return read_docx(file_path)
    elif file_type == FILE_TYPE_CSV:
        return read_csv(file_path)
    elif file_type == FILE_TYPE_XLSX:
        return read_xlsx(file_path)
    else:
        print(f"Unsupported file type for '{file_path}'")
    return None

def get_file_type(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    if file_extension == '.txt':
        return FILE_TYPE_TEXT
    elif file_extension == '.pdf':
        return FILE_TYPE_PDF
    elif file_extension == '.docx':
        return FILE_TYPE_DOCX
    elif file_extension == '.csv':
        return FILE_TYPE_CSV
    elif file_extension == '.xlsx':
        return FILE_TYPE_XLSX
    else:
        return None

def main():
    file_path = input("请输入文件路径: ")
    file_type = get_file_type(file_path)
    file_contents = None
    
    if file_type is None:
        print("不支持的文件类型。")
        return

    print(f"检测到的文件类型: {file_type}")
    
    if file_type == FILE_TYPE_TEXT:
        file_contents = read_text_file(file_path)
        print("选择操作:")
        print("1. 词频统计")
        print("2. 关键词提取")
        print("3. 搜索并替换")
        choice = input("请输入选项 (1/2/3): ")
        
        if choice == '1':
            word_frequency = count_word_frequency(file_contents)
            print(f"词频统计:\n{word_frequency}")
        elif choice == '2':
            num_keywords = int(input("要提取多少个关键词? "))
            keywords = extract_keywords(file_contents, num_keywords)
            print(f"关键词提取:\n{keywords}")
        elif choice == '3':
            search_term = input("请输入要搜索的词: ")
            replace_term = input("请输入替换后的词: ")
            updated_text = search_and_replace(file_contents, search_term, replace_term)
            print(f"更新后的文本内容:\n{updated_text}")
        else:
            print("无效的选项。")

    elif file_type == FILE_TYPE_PDF:
        file_contents = extract_text_from_pdf(file_path)
        print("选择操作:")
        print("1. 提取文本")
        print("2. 提取表格数据")
        choice = input("请输入选项 (1/2): ")
        
        if choice == '1':
            print(f"PDF文本内容:\n{file_contents}")
        elif choice == '2':
            table_data = extract_table_data_from_pdf(file_path)
            print(f"提取的表格数据:\n{table_data}")
        else:
            print("无效的选项。")

    elif file_type == FILE_TYPE_DOCX:
        file_contents = read_docx(file_path)
        print(f"DOCX文本内容:\n{file_contents}")
        update_content = input("请输入要替换的内容: ")
        update_docx(file_path, update_content)

    elif file_type == FILE_TYPE_CSV:
        file_contents = read_csv(file_path)
        clean_csv(file_path)
        print(f"CSV文件内容:\n{file_contents}")

    elif file_type == FILE_TYPE_XLSX:
        file_contents = read_xlsx(file_path)
        print(f"Excel文件内容:\n{file_contents}")
        export_data = input("请输入要导出的数据 (使用逗号分隔): ")
        export_to_excel([export_data.split(',')], 'output.xlsx')

    else:
        print(f"不支持的文件类型: '{file_path}'")

if __name__ == '__main__':
    main()
