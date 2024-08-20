import os
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
import pypandoc
from fpdf import FPDF
from jarvis import InputProcessor

def rotate_image(image):
    """逆时针旋转图像90度"""
    width, height = image.size
    if height > width:
        return image.rotate(90, expand=True)
    return image
    
def save_to_folder(output_file_path, output_folder):
    """保存文件到指定文件夹"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    return os.path.join(output_folder, os.path.basename(output_file_path))

def print_progress_bar(iteration, total, length=40):
    """打印进度条"""
    percent = ("{0:.1f}").format(100 * (iteration / float(total)))
    filled_length = int(length * iteration // total)
    bar = '█' * filled_length + '-' * (length - filled_length)
    sys.stdout.write(f'\r|{bar}| {percent}% 完成')
    sys.stdout.flush()

def convert_image_format(input_image_path, output_image_path):
    """将图片格式从一种转换为另一种"""
    image = Image.open(input_image_path)
    output_folder = os.path.dirname(output_image_path)
    output_file_path = save_to_folder(output_image_path, output_folder)

    # 检查输入和输出文件的扩展名
    input_ext = os.path.splitext(input_image_path)[1].lower()
    output_ext = os.path.splitext(output_image_path)[1].lower()

    # 如果需要转换格式，就保存为新的格式
    if input_ext != output_ext:
        image = rotate_image(image)
        image.save(output_file_path)
        print(f"已将 {input_image_path} 转换为 {output_file_path}")
    else:
        print(f"输入和输出格式相同，无需转换：{input_ext}")

def process_pdf(input_file_path, output_file_path, output_folder):
    reader = PdfReader(input_file_path)
    writer = PdfWriter()

    for page in reader.pages:
        if page.mediabox.height > page.mediabox.width:
            page.rotate_clockwise(90)
        
        if '/XObject' in page['/Resources']:
            xObject = page['/Resources']['/XObject'].get_object()
            for obj in xObject:
                if xObject[obj]['/Subtype'] == '/Image':
                    size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                    data = xObject[obj]._data

                    image = Image.open(BytesIO(data))
                    rotated_image = rotate_image(image)
                    # 替换PDF中的图像
                    img_buffer = BytesIO()
                    rotated_image.save(img_buffer, format=image.format)
                    img_buffer.seek(0)
                    page.images[img_index]['data'] = img_buffer.read()

        writer.add_page(page)
        print_progress_bar(i + 1, num_pages) # 更新进度条
    output_file_path = save_to_folder(output_file_path, output_folder)

    # 将旋转后的内容写入新的PDF
    with open(output_file_path, 'wb') as outfile:
        writer.write(outfile)

    print(f"处理并保存旋转PDF为 {output_file_path}")

def rotate_image_in_docx(image_path):
    """旋转DOCX文件中的图像，如果它是纵向的"""
    image = Image.open(image_path)
    return rotate_image(image)

def process_docx(input_file_path, output_file_path, output_folder):
    doc = Document(input_file_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                # 检测和旋转纵向文本
                run.text = run.text
                
    for shape in doc.inline_shapes:
        if shape.type == 3:  # 3代表图片
            # 如果需要，保存并旋转图像
            img_path = f"{shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name}.png"
            with open(img_path, 'wb') as img_file:
                img_file.write(shape._inline.graphic.graphicData.pic.blipFill.blip.embed._part.blob)

            rotated_image = rotate_image_in_docx(img_path)
            rotated_image.save(img_path)

            # 将图像替换为旋转后的版本
            shape._inline.graphic.graphicData.pic.blipFill.blip.embed._part = None
            shape._inline.graphic.graphicData.pic.blipFill.blip.embed = None
            shape._inline.graphic.graphicData.pic.blipFill.blip.embed._part = doc.add_picture(img_path, width=Inches(2))

        print_progress_bar(i + 1, num_shapes)  # 更新进度条
            
    output_file_path = save_to_folder(output_file_path, output_folder)
    doc.save(output_file_path)
    print(f"处理并保存旋转后的DOCX为 {output_file_path}")
    
def rotate_text(text):
    rotated_text = '\n'.join(text)
    return rotated_text 
       
def rotate_text_in_docx(input_file_path, output_file_path):
    doc = Document(input_file_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                # 检测和旋转纵向文本
                run.text = rotate_text(run.text)

    doc.save(output_file_path)
    print(f"只旋转文本并保存DOCX为 {output_file_path}") 
       
def images_to_pdf(image_paths, output_pdf_path):
    pdf = FPDF()
    
    for image_path in image_paths:
        image = Image.open(image_path)
        rotated_image = rotate_image(image)
        rotated_image_path = f"rotated_{os.path.basename(image_path)}"
        rotated_image.save(rotated_image_path)
        
        pdf.add_page(orientation="P" if rotated_image.width < rotated_image.height else "L")
        pdf.image(rotated_image_path, x=10, y=10, w=pdf.w - 20)
        
        os.remove(rotated_image_path)
        print_progress_bar(i + 1, num_images)

    pdf.output(output_pdf_path)
    print(f"Converted images to {output_pdf_path}")
    
def images_to_docx(image_paths, output_docx_path):
    doc = Document()

    for image_path in image_paths:
        image = Image.open(image_path)
        rotated_image = rotate_image(image)
        rotated_image_path = f"rotated_{os.path.basename(image_path)}"
        rotated_image.save(rotated_image_path)
        
        doc.add_picture(rotated_image_path, width=Inches(6))  # Adjust width as needed
        
        os.remove(rotated_image_path)
        print_progress_bar(i + 1, num_images)  
    doc.save(output_docx_path)
    print(f"Converted images to {output_docx_path}")
        
def convert_file(input_file_path, output_file_path, output_folder):
    input_ext = os.path.splitext(input_file_path)[1].lower()
    output_ext = os.path.splitext(output_file_path)[1].lower()

    if input_ext == '.pdf' and output_ext == '.docx':
        process_pdf(input_file_path, output_file_path, output_folder)
    elif input_ext == '.docx' and output_ext == '.pdf':
        process_docx(input_file_path, output_file_path, output_folder)
    elif input_ext == '.docx' and output_ext == '.txt':
        docx_to_txt(input_file_path, output_file_path)
    elif input_ext in ['.jpg', '.jpeg', '.png'] and output_ext == '.pdf':
        images_to_pdf([input_file_path], output_file_path, output_folder)
    elif input_ext in ['.jpg', '.jpeg', '.png'] and output_ext == '.docx':
        images_to_docx([input_file_path], output_file_path, output_folder)
    elif input_ext == '.txt' and output_ext == '.docx':
        txt_to_docx(input_file_path, output_file_path, output_folder)
    elif input_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff'] and output_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        convert_image_format(input_file_path, output_file_path)
    else:
        convert_with_pandoc(input_file_path, output_file_path, output_folder)

def docx_to_txt(input_file_path, output_file_path):
    doc = Document(input_file_path)
    with open(output_file_path, 'w', encoding='utf-8') as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + '\n')
    print(f"转换 {input_file_path} to {output_file_path}")

def txt_to_docx(input_file_path, output_file_path):
    doc = Document()
    with open(input_file_path, 'r', encoding='utf-8') as txt_file:
        for line in txt_file:
            doc.add_paragraph(line)
    doc.save(output_file_path)
    print(f"转换 {input_file_path} to {output_file_path}")

def convert_with_pandoc(input_file_path, output_file_path):
    pypandoc.convert_file(input_file_path, os.path.splitext(output_file_path)[1][1:], outputfile=output_file_path)
    print(f"转换 {input_file_path} to {output_file_path}")     
    
def file_converter():
    mode = 'voice'  # 默认语音输入模式
    inputProcessor = InputProcessor.InputProcessor()
    while True:
        if mode == 'voice':
            command = inputProcessor.process_voice_input()
        elif mode == 'text':
            command = inputProcessor.process_text_input()

        # 检查命令是否要切换模式
        if command == '1' and mode == 'voice':
            mode = 'text'
            continue
        elif command == '2' and mode == 'text':
            mode = 'voice'
            continue

        # 获取文件路径和转换后文件名
        input_file = input("请输入文件路径：")
        output_file = input("请输入要转换后的文件名：")
        output_folder = './my_package/downloaded'

        # 根据命令执行相应的文件处理
        if "处理pdf" in command:
            convert_file(input_file, output_file, output_folder)
        elif "处理docx" in command:
            convert_file(input_file, output_file, output_folder)
        elif "转换图片" in command:
            convert_file(input_file, output_file, output_folder)
        else:
            print("无法识别的命令，请重试。")
